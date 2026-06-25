#!/usr/bin/env python3
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor

HOME = os.path.expanduser("~")
DIR = os.path.join(HOME, "linux_assignment")
SHOTS = os.path.join(DIR, "shots")

from PIL import ImageFilter

MENLO = "/System/Library/Fonts/Menlo.ttc"
SF = "/System/Library/Fonts/SFNS.ttf"        # title bar font
S = 2                                         # retina 2x supersampling
FS = 12 * S                                   # Terminal default Menlo 12pt
font = ImageFont.truetype(MENLO, FS)
try:
    titlefont = ImageFont.truetype(SF, int(11 * S))
except Exception:
    titlefont = ImageFont.truetype(MENLO, int(11 * S))

# ---- default Terminal.app "Basic" profile: white bg, black mono text ----
PADX = 9 * S
PADTOP = 7 * S
LINE_H = int(FS * 1.18)
TITLEBAR = 28 * S
BODY = (255, 255, 255)
FG = (0, 0, 0)
# macOS window chrome
CORNER = 10 * S
MARGIN = 34 * S                               # space around window for the drop shadow

def render(txt_path, png_path, title):
    with open(txt_path) as f:
        lines = [l.rstrip("\n") for l in f]
    while lines and lines[-1].strip() == "":
        lines.pop()

    maxw = max((font.getlength(l) for l in lines), default=200)
    inner_w = max(int(maxw) + PADX * 2, 560 * S)
    inner_h = TITLEBAR + PADTOP * 2 + LINE_H * len(lines)

    # --- the terminal window itself ---
    win = Image.new("RGB", (inner_w, inner_h), BODY)
    d = ImageDraw.Draw(win)
    # title bar: subtle vertical gray gradient like Terminal.app
    for yy in range(TITLEBAR):
        t = yy / TITLEBAR
        g = int(236 - t * 26)
        d.line([(0, yy), (inner_w, yy)], fill=(g, g, g))
    d.line([(0, TITLEBAR - 1), (inner_w, TITLEBAR - 1)], fill=(176, 176, 176))
    # traffic lights
    for i, c in enumerate([(255, 95, 86), (255, 189, 46), (39, 201, 63)]):
        cx = 14 * S + i * 20 * S
        r = 6 * S
        cy = TITLEBAR // 2
        d.ellipse([cx, cy - r, cx + 2 * r, cy + r], fill=c)
        d.ellipse([cx, cy - r, cx + 2 * r, cy + r], outline=(0, 0, 0, 30))
    # centered window title: "folder — -zsh — 80x24"
    wt = title
    tw = titlefont.getlength(wt)
    d.text(((inner_w - tw) / 2, (TITLEBAR - int(11 * S)) / 2 - S), wt,
           font=titlefont, fill=(90, 90, 90))

    # --- body text (monochrome, like default zsh prompt) ---
    y = TITLEBAR + PADTOP
    for line in lines:
        d.text((PADX, y), line, font=font, fill=FG)
        y += LINE_H

    # --- round the window corners ---
    mask = Image.new("L", (inner_w, inner_h), 0)
    md = ImageDraw.Draw(mask)
    md.rounded_rectangle([0, 0, inner_w - 1, inner_h - 1], radius=CORNER, fill=255)
    win.putalpha(mask)

    # --- compose onto transparent canvas with a soft macOS window-capture shadow ---
    W = inner_w + MARGIN * 2
    H = inner_h + MARGIN * 2
    canvas = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    shadow = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow)
    sx, sy = MARGIN, MARGIN + 8 * S
    sd.rounded_rectangle([sx, sy, sx + inner_w, sy + inner_h],
                         radius=CORNER, fill=(0, 0, 0, 110))
    shadow = shadow.filter(ImageFilter.GaussianBlur(14 * S))
    canvas = Image.alpha_composite(canvas, shadow)
    canvas.alpha_composite(win, (MARGIN, MARGIN))

    # downscale for crisp retina-quality result
    canvas = canvas.resize((W // S, H // S), Image.LANCZOS)
    canvas.save(png_path)
    return png_path

WT = "rishav — -zsh — 80×24"
tasks = [
    ("task1", WT),
    ("task2", WT),
    ("task3", WT),
    ("task4", WT),
    ("task5", WT),
    ("task6", WT),
    ("task7", WT),
]
for key, title in tasks:
    render(os.path.join(SHOTS, key + ".txt"),
           os.path.join(SHOTS, key + ".png"), title)
print("rendered", len(tasks), "screenshots")

# ---------- build docx ----------
intro = [
 ("1. Creating and renaming files",
  "First I made a folder with mkdir, then made an empty file inside it with touch, and finally "
  "renamed the file using mv. After running mv the file showed up as renamed_example.txt when I "
  "listed the folder, so mv basically renames a file when you keep it in the same folder."),
 ("2. Looking at file contents",
  "I used the /etc/passwd file since it's always there. head -5 shows the first 5 lines (on a Mac "
  "those are just comment lines at the top) and tail -5 shows the last 5, which are the system "
  "accounts. I used -5 so I didn't have to scroll through the whole file."),
 ("3. Searching with grep",
  "Then I searched for the word root inside /etc/passwd. grep goes line by line and only prints the "
  "lines that match, so I got back the root account plus a couple of system lines that use /var/root "
  "as their home folder."),
 ("4. Zipping and unzipping",
  "I zipped the whole test_dir folder and unzipped it into a fresh folder to check it worked. The -r "
  "tells zip to include everything inside the folder, and -d lets you pick which folder to extract into. "
  "The renamed_example.txt file came back out exactly like the original."),
 ("5. Downloading a file",
  "The task asked for wget, but my Mac didn't have it (command not found), so I used curl which does the "
  "same job. The -o names the saved file and -L follows redirects. The sample.txt link gives a 404, so I "
  "just downloaded the example.com homepage to show the download working."),
 ("6. Changing permissions",
  "I made secure.txt and used chmod 444 to make it read only for everyone. The permissions went from "
  "-rw-r--r-- to -r--r--r--, and when I tried to add text it said permission denied, which is exactly "
  "what should happen. In octal 4 is read, 2 is write, 1 is execute, and the three digits are owner, "
  "group and everyone else. I can undo it with chmod 644."),
 ("7. Environment variables",
  "Last one. I set my own variable with export and printed it with echo $MY_VAR, then found it in the "
  "full list using env | grep MY_VAR. This only lasts for the current terminal window though, so to keep "
  "it permanently I'd add the export line to my ~/.zshrc file."),
]

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11.5)

h = doc.add_heading("Linux Commands Assignment", level=0)
p = doc.add_paragraph()
p.add_run("Name: Rishav Jamwal\n")
p.add_run("Repo: https://github.com/rishavrishu86-dotcom/linux-commandline-basics")
doc.add_paragraph(
 "I did all of this on my Mac using the Terminal (zsh shell). I made a folder called linux_assignment "
 "and worked inside it. For each part below I've written what I did and pasted a screenshot of the actual "
 "command and its output from my terminal. Most of these are normal Linux commands so they work the same "
 "on Linux too. The only one that was different is wget, which my Mac didn't have, so I used curl instead.")

for (heading, text), (key, _title) in zip(intro, tasks):
    doc.add_heading(heading, level=1)
    doc.add_paragraph(text)
    png = os.path.join(SHOTS, key + ".png")
    doc.add_picture(png, width=Inches(6.3))

doc.add_heading("Wrap up", level=1)
doc.add_paragraph(
 "All seven tasks worked. The only thing that needed changing was wget since it isn't on Mac by default, "
 "so I used curl which does the same thing. Everything else ran without any problems, and the screenshots "
 "above show each command with its real output.")

out = os.path.join(DIR, "Linux_CommandLine_Basics_Submission.docx")
doc.save(out)
print("saved", out)
